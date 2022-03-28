#导入Document用以打开docx文件
from docx import Document
#导入qn进行文字的样式处理
from docx.oxml.ns import qn
#导入Pt、Cm、Inches等计量单位对应磅、厘米、英寸
from docx.shared import Pt,Cm,Inches,RGBColor
import time
from docx import Document
import paramiko 
import glob
import os
import pandas as pd
import logging
from bottle import *
from secrets import *
import logging

HTML = HTML = """
<!DOCTYPE html>
<html>
<head>
<meta http-equiv="Content-Type" content="text/html; charset=utf-8" />
<title>自动生成安全巡检报告</title>
<link href="http://libs.baidu.com/bootstrap/3.0.3/css/bootstrap.min.css" rel="stylesheet" />
<script src="http://libs.baidu.com/jquery/2.1.4/jquery.min.js"></script>
<script src="http://libs.baidu.com/bootstrap/3.0.3/js/bootstrap.min.js"></script>
</head>
<body class="container">
    <nav class="navbar navbar-default">
        <div class="container-fluid">
          <div class="navbar-header">
            <a class="navbar-brand" href="#">自动生成安全巡检报告</a>
          </div>
          <div class="collapse navbar-collapse" id="bs-example-navbar-collapse-1">
            <ul class="nav navbar-nav">
              <li><a href="https://confluence.qingteng.cn/pages/viewpage.action?pageId=113143932">点这里查看更多使用技巧！！</a></li>
            </ul>
            <ul class="nav navbar-nav navbar-right">
               <li> <a>by北区交付中心</a></li>
            </ul>
          </div>
        </div>
      </nav>
      <div class="row">
        <div class="col-md-6">
            <p class="lead">
            <span class="glyphicon glyphicon-pushpin" aria-hidden="true"></span> 使用方式<br>
            将csv们打成一个zip或者tar，然后上传文件，填入信息，点执行 (请不要压缩文件夹，请选中多个文件框住，然后右键压缩)
            </p>
            <p class="lead">
            (需要漏洞检测-统计视图、3种风险、linux\windows主机详情、linux\windows弱密码  共 1+3+2+2 = 8张表格,如果某张表没有数据，就不放这张表)
            </p>
            <p class="lead">
            目前服务器是小水管，请不要反复点击提交，目前生成每份巡检需10秒。如果没有立即展示结果，建议保持页面不要关闭，等1~2分钟后回来看看
            </p>
        </div>
        <div class="col-md-6">
            <legend class="text-center">填写信息</legend>
            <form action="/upload" method="post" enctype="multipart/form-data">
                <div class="form-group">
                  <label for="客户全称">客户全称</label>
                  <input type='text' name='客户全称' id='客户全称' class='txt' placeholder='请输入客户全称: '/>
                </div>
                <div class="form-group">
                    <label for="客户简称">客户简称</label>
                    <input type='text' name='客户简称' id='客户简称' class='txt' placeholder='请输入客户简称: '/> 
                </div>
                <div class="form-group">
                    <label for="你的名字">你的名字</label>
                   <input type='text' name='你的名字' id='你的名字' class='txt' placeholder='请输入你的名字: '/>
                </div>
                <div class="form-group">
                    <label for="你的邮箱">你的邮箱</label>
                    <input type='text' name='你的邮箱' id='你的邮箱' class='txt' placeholder='请输入你的邮箱: '/>
                </div>
                <div class="form-group">
                  <label for="fileField">请选择压缩包文件</label>
                  <input type="file" name="fileField" class="file" id="fileField" size="28" onchange="document.getElementById('textfield').value=this.value" />
                  <p class="help-block">请选择zip或者tar文件</p>
                </div>
                <div class="checkbox">
                </div>
                <input type="submit" name="submit" class="btn" value="执行" onclick=""/>
              </form>
        </div>
      </div>
</body>
</html>
"""
 
base_path = os.path.dirname(os.path.realpath(__file__))  # 获取脚本路径
upload_path = os.path.join(base_path, 'upload')   # 上传文件目录

if not os.path.exists(upload_path):
	os.makedirs(upload_path)
 
 
@route('/', method='GET')
@route('/upload', method='GET')
@route('/index.html', method='GET')
@route('/upload.html', method='GET')
def index():
	"""显示上传页"""
	return HTML
 
 
@route('/upload', method='POST')
def do_upload():
	"""处理上传文件"""
	filedata = request.files.get('fileField')
	
	if filedata.file:
		path = str(time.time())+token_hex(16)
		if ".zip" in filedata.filename:
			filedata.filename = path
			file_name = os.path.join(upload_path, filedata.filename)
			try:
				filedata.save(file_name+".zip")  # 上传文件写入
				os.system(f'unzip -O GBK  {file_name}.zip -d {file_name}')
				# ~ print (f'unzip -O GBK  {file_name}.zip -d {file_name}')
			except IOError:
				return '上传文件失败'
		elif ".tar" in filedata.filename:
			filedata.filename = path
			file_name = os.path.join(upload_path, filedata.filename)
			try:
				filedata.save(file_name+".tar")  # 上传文件写入
				os.makedirs(file_name)
				os.system(f'tar -xvf {file_name}.tar -C {file_name}')
				# ~ print (f'tar -xvf {file_name}.tar -C {file_name}/')
			except IOError:
				return '上传文件失败'
		else:
			return '上传文件失败, 请以压缩包方式提交文件'
		
		
		df_count = get_count(file_name)
		
		replace_dict = {
		"$CName":str(request.forms.get('客户简称')),
		"$Date": time.strftime('%Y-%m-%d',time.localtime(time.time())),
		"$CompanyName":str(request.forms.get('客户全称')),
		"$count_all":str(df_count["在线"]+df_count["离线"]),
		"$count_online":str(df_count["在线"]),
		"$count_offline":str(df_count["离线"]),
		"$count_linux":str(df_count["linux总数"]),
		"$count_windows":str(df_count["windows总数"]),
		"$online_rate":str(100*df_count["在线"]/(df_count["在线"]+df_count["离线"]))[0:5],
		"$user_name":str(request.forms.get('你的名字')),
		"$user_email_name":str(request.forms.get('你的邮箱'))
		}

		doc = getdoc("./docx/template2.docx",replace_dict, file_name)
		doc.save(file_name+"/"+str(request.forms.get('客户简称'))+"安全巡检报告.docx")
		
			
			

		return '已完成巡检, 报告下载连接: '+"192.168.192.77:16932/"+path+"/"+str(request.forms.get('客户简称'))+"安全巡检报告.docx"
        
        
	else:
		return '上传文件失败'
 
@route('/favicon.ico', method='GET')
def server_static():
	"""处理网站图标文件, 找个图标文件放在脚本目录里"""
	return static_file('favicon.ico', root=base_path)
 
 
@error(404)
def error404(error):
	"""处理错误信息"""
	return '404 发生页面错误, 未找到内容'
 




def CSVRead(filename):
	"""
	读取CSV文件
	input：CSV文件名
	output： 对应的DataFrame（读取失败则返回0）
	"""
	try:
		file = pd.read_csv(filename)
		df = pd.DataFrame(file)
		return(df)
	except:
		try:
			file = pd.read_csv(filename, encoding = "gb18030")
			df = pd.DataFrame(file)
			return(df)
		except (Exception, BaseException) as e:
			logging.error(e)
			logging.error(f'文件取失败，请确认文件是否存在！')
			print(e)
			print(f'文件读取失败，请确认文件是否存在！')
			return(0)


def get_OS(file_name):
	"""
	读取本地的主机列表csv，统计各操作系统数量	
	output: 统计后的数据以DataFrame格式返回
	"""
	for file in glob.glob(os.path.join(file_name, 'Linux-托管主机列表*.csv')):
		df1 = CSVRead(file)
	for file in glob.glob(os.path.join(file_name, 'Win-主机基本信息列表*.csv')):
		df2 = CSVRead(file)

	series1 = df1['操作系统'].value_counts()
	series2 = df2['操作系统'].value_counts()
	df = pd.concat([series1, series2], axis=0)
	return(df)
	
def get_linux_weak_password(app,file_name):
	"""
	读取本地的弱密码列表csv，筛选特定应用的弱密码以DataFrame格式返回 	
	input: 应用名
	output: 筛选后的数据以DataFrame格式返回
	"""
	result={}
	index = 0
	df1 = pd.DataFrame(columns=['应用']) 
	for file in glob.glob(os.path.join(file_name, 'Linux弱密码_全部风险项列表*.csv')):
		df1 = CSVRead(file)
	for index, row in df1.iterrows():
		if row['应用'] == str(app):
			result[index]=row
			# ~ print(result)
			# ~ print("======================================")
			index = index+1
	df = pd.DataFrame.from_dict(result, orient='index')
	return(df)


def get_windows_weak_password(file_name):
	"""
	读取本地的windows弱密码列表csv，以DataFrame格式返回 	
	output: 数据以DataFrame格式返回
	"""
	result={}
	index = 0
	df1 = pd.DataFrame(columns=['A', 'B', 'C', 'D']) 
	for file in glob.glob(os.path.join(file_name, 'windows弱密码*.csv')):
		df1 = CSVRead(file)
	return(df1)


def get_count(file_name):
	"""
	读取本地的主机列表csv，统计在线，离线，总计agent数量	
	output: 统计后的数据以字典格式返回
	"""
	#print(file_name)
	for file in glob.glob(os.path.join(file_name, 'Linux-托管主机列表*.csv')):
		df1 = CSVRead(file)
	for file in glob.glob(os.path.join(file_name, 'Win-主机基本信息列表*.csv')):
		df2 = CSVRead(file)
	
	js={"离线":"","在线":"","linux总数":"","windows总数":""}
	series1 = df1['主机状态'].value_counts()
	series2 = df2['主机状态'].value_counts()
	if "离线" not in series1:
		series1["离线"]=0
	if "在线" not in series1:
		series1["在线"]=0
	if "在线" not in series2:
		series2["在线"]=0
	if "离线" not in series2:
		series2["离线"]=0
	js["离线"]=series1["离线"]+series2["离线"]
	js["在线"]=series1["在线"]+series2["在线"]
	js["linux总数"]=series1["在线"]+series1["离线"]
	js["windows总数"]=series2["在线"]+series2["离线"]

	return(js)


def check_and_change(document, replace_dict):
	"""
	docx全文替换模块，依据全局变量replace_dict对文档中的关键字进行替换（保留格式不变）
	input：文档对象,字典格式的替换列表
	output： 替换后的文档对象
	"""
	for para in document.paragraphs:	#遍历文档中的所有段落，进行关键字替换
		for key, value in replace_dict.items():
			if key in para.text:
				logging.info(key+"->"+value)
				style = para.style		#由于替换时会覆盖原格式，需要保存原格式手动恢复
				para.text = para.text.replace(key, value)
				para.style = style
	for tab in document.tables: 	#遍历文档中的所有表格，进行关键字替换
		for w in tab.rows:
			for cell in w.cells:
				for para in cell.paragraphs:
					for key, value in replace_dict.items():
						if key in para.text:
							logging.info(key+"->"+value)
							style = para.style
							para.text = para.text.replace(key, value)
							para.style = style
	return document

def check_and_change2(document, replace_dict):
	"""
	docx全文替换模块，依据全局变量replace_dict对文档中的整段替换进行替换（保留格式不变）
	input：文档对象,字典格式的替换列表
	output： 替换后的文档对象
	"""
	for para in document.paragraphs:	#遍历文档中的所有段落，进行关键字替换
		for key, value in replace_dict.items():
			if key in para.text:
				logging.info(key+"->"+value)
				style = para.style		#由于替换时会覆盖原格式，需要保存原格式手动恢复
				para.text = para.text.replace(key, value)
				para.style = style
	for tab in document.tables: 	#遍历文档中的所有表格，进行关键字替换
		for w in tab.rows:
			for cell in w.cells:
				for para in cell.paragraphs:
					for key, value in replace_dict.items():
						if str(key) == str(para.text):
							logging.info(key+"->"+value)
							style = para.style
							para.text = para.text.replace(key, value)
							para.style = style
	return document
	
def remove_paragraphs(document, remove_list):
	"""
	段落删除模块，如果模块包含指定关键字，就删除它
	input：文档对象, list格式的关键字列表
	output： 新表格对象
	"""
	list2 = []
	for para in document.paragraphs:	#遍历文档中的所有段落，存在关键字就删除
		for key in remove_list:
			if key in para.text:
				list2.append(para)
	for x in list2:
		x._element.getparent().remove(x._element)
	return(document)
	
def remove_table_paragraphs(document, remove_list):
	"""
	段落删除模块，如果表格的段落包含指定关键字，就删除它
	input：文档对象, list格式的关键字列表
	output： 新文档对象
	"""
	list2 = []
	txt = ""
	for tab in document.tables: 	#遍历文档中的所有表格的所有段落，存在关键字就删除
		for w in tab.rows:
			for cell in w.cells:
				for para in cell.paragraphs:
					for key in remove_list:
						if key in para.text:
							for txt1 in para.text.split('\n'):	
								if key not in txt1:
									txt = txt+txt1
							style = para.style
							para.text = txt
							para.style = style
							txt = ""
	return(document)


def add_weak_password(table, app, file_name):
	"""
	标准弱密码table填写模块
	input：表格对象, app
	output： 若有对应应用弱密码返回新表格对象,如果有对没应应用弱密码，返回0
	"""
	logging.info("已捕获$"+app+"弱密码问题列表")
	logging.info("开始为表格填入内容")
	cell = table.rows[0].cells[0]
	style = cell.paragraphs[0].style
	cell.paragraphs[0].text = "主机IP"
	cell.paragraphs[0].style = style 
	df2 = get_linux_weak_password(app,file_name)
	if df2.empty: #如果没有弱口令，删除这张表
		logging.info("没有"+app+"弱口令")
		return(0)
	count = 1
	for index, row1 in df2.iterrows():
		if count < 50:
			list1 =[row1["主机IP"],row1["账号名"],row1["密码值"],row1["弱密码类型"],row1["监听端口"]]
			table = table_add_a_line(table, list1)
			count = count+1
		elif count == 50:
			list1 =["...完整弱密码列表请见：附件-弱密码列表","","","","",]
			table = table_add_a_line(table[q], list1)
			break
	table._tbl.remove(table.rows[len(table.rows)-1]._tr)#删除最后一行（最后一行是空的）
	return(table)
	
def table_add_a_line(table, js):
	"""
	向docx表格插入一行数据
	input：表格对象, 字典格式的要插入的数据
	output： 新表格对象
	"""
	table.add_row()
	for x in range(0,len(js)):
		style1 = table.cell(len(table.rows)-3,x).paragraphs[0].style
		style2 = table.cell(len(table.rows)-3,x).paragraphs[0].paragraph_format.alignment
		table.cell(len(table.rows)-2,x).text= str(js[x])
		table.cell(len(table.rows)-2,x).paragraphs[0].style = style1
		table.cell(len(table.rows)-2,x).paragraphs[0].paragraph_format.alignment = style2
		
		#字体颜色设置一直不成功，最后再弄它
		# ~ if table.cell(len(table.rows)-2,x).paragraphs[0].text == "危急":
			# ~ table.cell(len(table.rows)-3,x).paragraphs[0].runs[0].style.font.color.rgb = RGBColor(187, 0, 85)
			# ~ print(table.cell(len(table.rows)-2,x).paragraphs[0].runs[0].style.font.color.rgb)
			# ~ table.cell(len(table.rows)-2,x).paragraphs[0].runs[0].style.font.color.rgb = RGBColor(187, 0, 85)
			# ~ print(table.cell(len(table.rows)-2,x).paragraphs[0].runs[0].style.font.color.rgb)
		# ~ if table.cell(len(table.rows)-3,x).paragraphs[0].text == "高危":
			# ~ table.cell(len(table.rows)-3,x).paragraphs[0].style.font.color.rgb = RGBColor(238, 0, 0)
	return(table)
	
	
	
	
def getdoc(docx_dir, replace_dict, file_name):
	logging.info("=====================================================================")
	logging.info("开始巡检，本次巡检的基本信息")
	logging.info(replace_dict)
	if os.path.exists(file_name+"/template2.docx"): #如果上传了模板，使用上传上来的模板
		doc = Document(file_name+"/template2.docx")
		logging.info("使用自定义模板")
	else:
		doc = Document(docx_dir)
	doc = check_and_change(doc, replace_dict)
	#文字处理获取段落，形成列表
	pg = doc.paragraphs
	table = doc.tables
	no_data_tables = [] 
	no_data_paragraphs = []
	# ~ #由于段落是列表，遍历，并对每一段落进行处理实现全文修改
	has_text = 0
	count_vul = 0
	count_app_vul = 0
	count_os_vul = 0
	count_account_vul = 0
	for q in range(0,len(table)):
		for i, row in enumerate(table[q].rows):
			# ~ print(q,"",i)
			for w in range(0,len(row.cells)):
				cell = row.cells[w]
				# ~ print(cell.text)
				# 生成操作系统分布情况
				if (cell.text == "$操作系统分布表"):
					# ~ logging.info("已捕获$操作系统分布表")
					# ~ logging.info("开始为表格填入内容")
					style = cell.paragraphs[0].style
					cell.paragraphs[0].text = "系统版本"
					cell.paragraphs[0].style = style 
					df = get_OS(file_name)
					for i, x in df.items():
						list1=[i,x]
						table[q] = table_add_a_line(table[q], list1)
					table[q]._tbl.remove(table[q].rows[len(table[q].rows)-1]._tr)
					
					
				# 生成弱密码列表（每种弱密码只显示50个，若超过50个，在第51行显示：...完整弱密码列表请见：附件-弱密码列表）
				if (cell.text == "$MySQL弱密码问题列表"):
					logging.info("已捕获$MySQL弱密码问题列表")
					logging.info("开始为表格填入内容")
					style = cell.paragraphs[0].style
					cell.paragraphs[0].text = "主机IP"
					cell.paragraphs[0].style = style 
					df2 = get_linux_weak_password("MySQL",file_name)
					if df2.empty: #如果没有mysql空口令，删除这张表
						logging.info("没有mysql弱口令")
						no_data_tables.append(table[q])#将没有数据的表存入数组，后续一并删除
						no_data_paragraphs.append("$MySQL弱密码问题列表")#对应的说明字段也存入数组，后续一并删除
						break
					mysql_count = 1
					for index, row1 in df2.iterrows():
						if mysql_count < 50:
							list1 =[row1["主机IP"],row1["监听端口"],row1["账号名"],row1["密码值"],row1["弱密码类型"],]
							table[q] = table_add_a_line(table[q], list1)
							mysql_count = mysql_count+1
						elif mysql_count == 50:
							list1 =["...完整弱密码列表请见：附件-弱密码列表","","","","",]
							table[q] = table_add_a_line(table[q], list1)
							break
					table[q]._tbl.remove(table[q].rows[len(table[q].rows)-1]._tr)#删除最后一行（最后一行是空的）
					
				if (cell.text == "$SSH弱密码问题列表"):
					logging.info("已捕获$SSH弱密码问题列表")
					logging.info("开始为表格填入内容")
					style = cell.paragraphs[0].style
					cell.paragraphs[0].text = "主机IP"
					cell.paragraphs[0].style = style 
					df2 = get_linux_weak_password("SSH",file_name)					
					if df2.empty: #如果没有SSH空口令，删除这张表
						logging.info("没有SSH弱口令")
						no_data_tables.append(table[q])#将没有数据的表存入数组，后续一并删除
						no_data_paragraphs.append("$SSH弱密码问题列表")#对应的说明字段也存入数组，后续一并删除
						break
					SSH_count = 1
					for index, row1 in df2.iterrows():
						# ~ logging.info("get row: ",index)
						if SSH_count < 50:
							list1 =[row1["主机IP"],row1["账号名"],row1["SSH登录方式"],row1["密码值"],row1["弱密码类型"],]
							table[q] = table_add_a_line(table[q], list1)
							SSH_count = SSH_count+1
						elif SSH_count == 50:
							list1 =["...完整弱密码列表请见：附件-弱密码列表","","","","",]
							table[q] = table_add_a_line(table[q], list1)
							break
					table[q]._tbl.remove(table[q].rows[len(table[q].rows)-1]._tr)#删除最后一行（最后一行是空的）
				
				if (cell.text == "$Redis弱密码问题列表"):
					logging.info("已捕获$Redis弱密码问题列表")
					logging.info("开始为表格填入内容")
					style = cell.paragraphs[0].style
					cell.paragraphs[0].text = "主机IP"
					cell.paragraphs[0].style = style 
					df2 = get_linux_weak_password("Redis",file_name)
					if df2.empty: #如果没有SSH空口令，删除这张表
						logging.info("没有Redis弱口令")
						no_data_tables.append(table[q])#将没有数据的表存入数组，后续一并删除
						no_data_paragraphs.append("$Redis弱密码问题列表")#对应的说明字段也存入数组，后续一并删除
						break
					Redis_count = 1
					for index, row1 in df2.iterrows():
						if Redis_count < 50:
							list1 =[row1["主机IP"],row1["监听端口"],row1["绑定IP"],row1["是否root权限"],row1["弱密码类型"],]
							table[q] = table_add_a_line(table[q], list1)
							Redis_count = Redis_count+1
						elif Redis_count == 50:
							list1 =["...完整弱密码列表请见：附件-弱密码列表","","","","",]
							table[q] = table_add_a_line(table[q], list1)
							break
					table[q]._tbl.remove(table[q].rows[len(table[q].rows)-1]._tr)#删除最后一行（最后一行是空的）
					
				if (cell.text == "$Tomcat弱密码问题列表"):
					logging.info("已捕获$Tomcat弱密码问题列表")
					logging.info("开始为表格填入内容")
					style = cell.paragraphs[0].style
					cell.paragraphs[0].text = "主机IP"
					cell.paragraphs[0].style = style 
					df2 = get_linux_weak_password("Tomcat",file_name)
					if df2.empty: #如果没有SSH空口令，删除这张表
						logging.info("没有Tomcat弱口令")
						no_data_tables.append(table[q])#将没有数据的表存入数组，后续一并删除
						no_data_paragraphs.append("$Tomcat弱密码问题列表")#对应的说明字段也存入数组，后续一并删除
						break
					Tomcat_count = 0
					for index, row1 in df2.iterrows():
						if Tomcat_count < 50:
							list1 =[row1["主机IP"],row1["账号状态"],row1["账号名"],row1["密码值"],row1["弱密码类型"],row1["监听端口"]]
							table[q] = table_add_a_line(table[q], list1)
							Tomcat_count = Tomcat_count+1
						elif Tomcat_count == 50:
							list1 =["...完整弱密码列表请见：附件-弱密码列表","","","","",]
							table[q] = table_add_a_line(table[q], list1)
							break
					table[q]._tbl.remove(table[q].rows[len(table[q].rows)-1]._tr)#删除最后一行（最后一行是空的）		
					
				#以上为输出字段非标准的弱密码项，以下为输出字段标准的弱密码项==================
				if (cell.text == "$Rsync弱密码问题列表"):
					t = add_weak_password(table[q], "Rsync", file_name)
					if t != 0:
						 table[q] = t
					elif t == 0:
						no_data_tables.append(table[q])#将没有数据的表存入数组，后续一并删除
						no_data_paragraphs.append("$Rsync弱密码问题列表")#对应的说明字段也存入数组，后续一并删除					 
					
				if (cell.text == "$Weblogic弱密码问题列表"):
					t = add_weak_password(table[q], "Weblogic", file_name)
					if t != 0:
						table[q] = t
					elif t == 0:
						no_data_tables.append(table[q])#将没有数据的表存入数组，后续一并删除
						no_data_paragraphs.append("$Weblogic弱密码问题列表")#对应的说明字段也存入数组，后续一并删除
					
				if (cell.text == "$Jenkins弱密码问题列表"):
					t = add_weak_password(table[q], "Jenkins", file_name)
					if t != 0:
						table[q] = t
					elif t == 0:
						no_data_tables.append(table[q])#将没有数据的表存入数组，后续一并删除
						no_data_paragraphs.append("$Jenkins弱密码问题列表")#对应的说明字段也存入数组，后续一并删除		
					
				if (cell.text == "$OpenLDAP弱密码问题列表"):
					t = add_weak_password(table[q], "OpenLDAP", file_name)
					if t != 0:
						table[q] = t
					elif t == 0:
						no_data_tables.append(table[q])#将没有数据的表存入数组，后续一并删除
						no_data_paragraphs.append("$OpenLDAP弱密码问题列表")#对应的说明字段也存入数组，后续一并删除
						
				if (cell.text == "$InfluxDB弱密码问题列表"):
					t = add_weak_password(table[q], "InfluxDB", file_name)
					if t != 0:
						table[q] = t
					elif t == 0:
						no_data_tables.append(table[q])#将没有数据的表存入数组，后续一并删除
						no_data_paragraphs.append("$InfluxDB弱密码问题列表")#对应的说明字段也存入数组，后续一并删除
						
				if (cell.text == "$ProFTPD弱密码问题列表"):
					t = add_weak_password(table[q], "ProFTPD", file_name)
					if t != 0:
						table[q] = t
					elif t == 0:
						no_data_tables.append(table[q])#将没有数据的表存入数组，后续一并删除
						no_data_paragraphs.append("$ProFTPD弱密码问题列表")#对应的说明字段也存入数组，后续一并删除
				
				if (cell.text == "$SVN弱密码问题列表"):
					t = add_weak_password(table[q], "SVN", file_name)
					if t != 0:
						table[q] = t
					elif t == 0:
						no_data_tables.append(table[q])#将没有数据的表存入数组，后续一并删除
						no_data_paragraphs.append("$SVN弱密码问题列表")#对应的说明字段也存入数组，后续一并删除
				
				if (cell.text == "$PPTP弱密码问题列表"):
					t = add_weak_password(table[q], "PPTP", file_name)
					if t != 0:
						table[q] = t
					elif t == 0:
						no_data_tables.append(table[q])#将没有数据的表存入数组，后续一并删除
						no_data_paragraphs.append("$PPTP弱密码问题列表")#对应的说明字段也存入数组，后续一并删除
				
				if (cell.text == "$VNC弱密码问题列表"):
					t = add_weak_password(table[q], "VNC", file_name)
					if t != 0:
						table[q] = t
					elif t == 0:
						no_data_tables.append(table[q])#将没有数据的表存入数组，后续一并删除
						no_data_paragraphs.append("$VNC弱密码问题列表")#对应的说明字段也存入数组，后续一并删除
						
				if (cell.text == "$OpenVPN弱密码问题列表"):
					t = add_weak_password(table[q], "OpenVPN", file_name)
					if t != 0:
						table[q] = t
					elif t == 0:
						no_data_tables.append(table[q])#将没有数据的表存入数组，后续一并删除
						no_data_paragraphs.append("$OpenVPN弱密码问题列表")#对应的说明字段也存入数组，后续一并删除
						
				if (cell.text == "$vsftpd弱密码问题列表"):
					t = add_weak_password(table[q], "vsftpd", file_name)
					if t != 0:
						table[q] = t
					elif t == 0:
						no_data_tables.append(table[q])#将没有数据的表存入数组，后续一并删除
						no_data_paragraphs.append("$vsftpd弱密码问题列表")#对应的说明字段也存入数组，后续一并删除
						
				if (cell.text == "$windows系统账户弱密码问题列表"):
					logging.info("已捕获$windows系统账户弱密码问题列表")
					logging.info("开始为表格填入内容")
					style = cell.paragraphs[0].style
					cell.paragraphs[0].text = "主机IP"
					cell.paragraphs[0].style = style 
					df2 = get_windows_weak_password(file_name)
					if df2.empty: #如果没有windows空口令，删除这张表
						logging.info("没有windows弱口令")
						no_data_tables.append(table[q])#将没有数据的表存入数组，后续一并删除
						no_data_paragraphs.append("$windows系统账户弱密码问题列表")#对应的说明字段也存入数组，后续一并删除
						break
					for index, row1 in df2.iterrows():
						if index < 50:
							list1 =[row1["主机IP"],row1["弱密码帐户"],row1["账号状态"],row1["密码值"],row1["弱密码类型"]]
							table[q] = table_add_a_line(table[q], list1)
						elif index == 50:
							list1 =["...完整弱密码列表请见：附件-弱密码列表-windows","","","","",]
							table[q] = table_add_a_line(table[q], list1)
							break
					table[q]._tbl.remove(table[q].rows[len(table[q].rows)-1]._tr)#删除最后一行（最后一行是空的）
					
				# 生成漏洞情况,目前固定输出10个最危险的漏洞
				if (cell.text == "$Linux漏洞风险"):
					logging.info("已捕获$Linux漏洞风险")
					logging.info("开始为表格填入内容")
					style = cell.paragraphs[0].style
					cell.paragraphs[0].text = "漏洞名称"
					cell.paragraphs[0].style = style 
					count_vul = 0 #统计已写入的漏洞数量数量
					vul_list = []
					if os.path.exists(file_name+"/必修漏洞列表.csv"): #如果上传了模板，使用上传上来的模板
						df1 = CSVRead(file_name+"/必修漏洞列表.csv")
						logging.info("使用自定义 必修漏洞列表")
					else:
						df1 = CSVRead("./csv/必修漏洞列表.csv")
					df2 = pd.DataFrame(columns=['A', 'B', 'C', 'D']) 
					for file in glob.glob(os.path.join(file_name, 'Linux漏洞检测_风险视图统计列表_*.csv')):
						df2 = CSVRead(file)
					for index, row_important_vul in df1.iterrows():
						for index2, row_vul_list in df2.iterrows():
							if row_important_vul["漏洞名称"] == row_vul_list["漏洞名称"]:
								list1 =[row_vul_list["漏洞名称"],str(row_vul_list["影响主机数"])+"台",row_vul_list["漏洞类型"],row_vul_list["危险程度"]]
								table[q] = table_add_a_line(table[q], list1)
								list2 = [row_vul_list["漏洞名称"],str(row_vul_list["影响主机数"])+"台",row_vul_list["漏洞描述"]+"\r【修复建议】"+row_vul_list["修复建议"]+"\r【修复影响】"+row_vul_list["修复影响"]]
								table[q+1] = table_add_a_line(table[q+1], list2)
								vul_list.append(row_vul_list["漏洞名称"])
								count_vul = count_vul+1
					if count_vul < 10: #如果最高优先级漏洞全都匹配完，输出的漏洞不够10个，就在剩下的漏洞中选择可远程利用、危急或高危，存在exp的漏洞
						if len(df2) > 20:
							for index2, row_vul_list in df2.iterrows():
								if row_vul_list["漏洞名称"] not in vul_list:
									if row_vul_list["远程利用"] == "是" and row_vul_list["存在EXP"] == "是":
										if row_vul_list["危险程度"] == "危急" or row_vul_list["危险程度"] == "高危":
											list1 =[row_vul_list["漏洞名称"],str(row_vul_list["影响主机数"])+"台",row_vul_list["漏洞类型"],row_vul_list["危险程度"]]
											table[q] = table_add_a_line(table[q], list1)
											list2 = [row_vul_list["漏洞名称"],str(row_vul_list["影响主机数"])+"台",row_vul_list["漏洞描述"]+"\r【修复建议】"+row_vul_list["修复建议"]+"\r【修复影响】"+row_vul_list["修复影响"]]
											table[q+1] = table_add_a_line(table[q+1], list2)
											vul_list.append(row_vul_list["漏洞名称"])
											count_vul = count_vul+1
											if count_vul > 9:
												break
					table[q]._tbl.remove(table[q].rows[len(table[q].rows)-1]._tr)#删除最后一行（最后一行是空的）
					table[q+1]._tbl.remove(table[q+1].rows[len(table[q+1].rows)-1]._tr)#删除最后一行（最后一行是空的）
												
				# 生成应用风险情况,目前固定输出10个最危险的漏洞
				if (cell.text == "$应用风险检查情况"):
					logging.info("已捕获$应用风险检查情况")
					logging.info("开始为表格填入内容")
					style = cell.paragraphs[0].style
					cell.paragraphs[0].text = "危险程度"
					cell.paragraphs[0].style = style 
					count_app_vul = 0 #统计已写入的漏洞数量数量
					vul_list = []
					df1 = pd.DataFrame(columns=['A', 'B', 'C', 'D']) 
					for file in glob.glob(os.path.join(file_name, 'Linux应用风险_风险视图列表*.csv')):
						df1 = CSVRead(file)
					if df1.empty: #如果没有应用风险,删除这张表
						logging.info("没有应用风险")
						no_data_tables.append(table[q])#将没有数据的表存入数组，后续一并删除
						no_data_tables.append(table[q+1])
						no_data_paragraphs.append("$应用风险检查情况")#对应的说明字段也存入数组，后续一并删除
						break
					#选择危急或高危的漏洞
					for index, row_vul_list in df1.iterrows():
						if row_vul_list["风险项名"] not in vul_list:
							if row_vul_list["危险程度"] == "危急":
								list1 =[row_vul_list["危险程度"], row_vul_list["风险项名"], str(df1.value_counts("风险项名")[row_vul_list["风险项名"]])+"台"]
								table[q] = table_add_a_line(table[q], list1)
								list2 = [row_vul_list["风险项名"],str(df1.value_counts("风险项名")[row_vul_list["风险项名"]])+"台",row_vul_list["风险描述"]+"\r【修复建议】"+row_vul_list["修复建议"]+"\r\r【修复影响】"+row_vul_list["修复影响"]]
								table[q+1] = table_add_a_line(table[q+1], list2)
								count_app_vul = count_app_vul+1
								vul_list.append(row_vul_list["风险项名"])
								if count_app_vul > 9:
									break
							if row_vul_list["危险程度"] == "高危":
								list1 =[row_vul_list["危险程度"], row_vul_list["风险项名"], str(df1.value_counts("风险项名")[row_vul_list["风险项名"]])+"台"]
								table[q] = table_add_a_line(table[q], list1)
								list2 = [row_vul_list["风险项名"],str(df1.value_counts("风险项名")[row_vul_list["风险项名"]])+"台",row_vul_list["风险描述"]+"\r【修复建议】"+row_vul_list["修复建议"]+"\r\r【修复影响】"+row_vul_list["修复影响"]]
								table[q+1] = table_add_a_line(table[q+1], list2)
								count_app_vul = count_app_vul+1
								vul_list.append(row_vul_list["风险项名"])
								if count_app_vul > 9:
									break
					table[q]._tbl.remove(table[q].rows[len(table[q].rows)-1]._tr)#删除最后一行（最后一行是空的）
					table[q+1]._tbl.remove(table[q+1].rows[len(table[q+1].rows)-1]._tr)#删除最后一行（最后一行是空的）
												
				# 生成系统风险情况,目前固定输出10个最危险的漏洞
				if (cell.text == "$系统风险检查情况"):
					logging.info("已捕获$系统风险检查情况")
					logging.info("开始为表格填入内容")
					style = cell.paragraphs[0].style
					cell.paragraphs[0].text = "危险程度"
					cell.paragraphs[0].style = style 
					count_os_vul = 0 #统计已写入的漏洞数量数量
					vul_list = []
					df1 = pd.DataFrame(columns=['A', 'B', 'C', 'D']) 
					for file in glob.glob(os.path.join(file_name, 'Linux系统风险_风险视图列表*.csv')):
						df1 = CSVRead(file)
					if df1.empty: #如果没有系统风险,删除这张表
						logging.info("没有系统风险")
						no_data_tables.append(table[q])#将没有数据的表存入数组，后续一并删除
						no_data_tables.append(table[q+1])
						no_data_paragraphs.append("$系统风险检查情况")#对应的说明字段也存入数组，后续一并删除
						break
					#选择危急或高危的漏洞
					for index, row_vul_list in df1.iterrows():
						if row_vul_list["风险项名"] not in vul_list:
							if row_vul_list["危险程度"] == "危急":
								list1 =[row_vul_list["危险程度"], row_vul_list["风险项名"], str(df1.value_counts("风险项名")[row_vul_list["风险项名"]])+"台"]
								table[q] = table_add_a_line(table[q], list1)
								list2 = [row_vul_list["风险项名"],str(df1.value_counts("风险项名")[row_vul_list["风险项名"]])+"台",row_vul_list["风险描述"]+"\r【修复建议】"+row_vul_list["修复建议"]+"\r\r【修复影响】"+row_vul_list["修复影响"]]
								table[q+1] = table_add_a_line(table[q+1], list2)
								count_os_vul = count_os_vul+1
								vul_list.append(row_vul_list["风险项名"])
								if count_os_vul > 9:
									break
							if row_vul_list["危险程度"] == "高危":
								list1 =[row_vul_list["危险程度"], row_vul_list["风险项名"], str(df1.value_counts("风险项名")[row_vul_list["风险项名"]])+"台"]
								table[q] = table_add_a_line(table[q], list1)
								list2 = [row_vul_list["风险项名"],str(df1.value_counts("风险项名")[row_vul_list["风险项名"]])+"台",row_vul_list["风险描述"]+"\r【修复建议】"+row_vul_list["修复建议"]+"\r\r【修复影响】"+row_vul_list["修复影响"]]
								table[q+1] = table_add_a_line(table[q+1], list2)
								count_os_vul = count_os_vul+1
								vul_list.append(row_vul_list["风险项名"])
								if count_os_vul > 9:
									break
					# ~ table[q]._tbl.remove(table[q].rows[len(table[q].rows)-1]._tr)#删除最后一行（最后一行是空的）
					table[q+1]._tbl.remove(table[q+1].rows[len(table[q+1].rows)-1]._tr)#删除最后一行（最后一行是空的）

				# 生成账号风险情况,目前固定输出10个最危险的漏洞
				if (cell.text == "$账号风险检查情况"):
					logging.info("已捕获$账号风险检查情况")
					logging.info("开始为表格填入内容")
					style = cell.paragraphs[0].style
					cell.paragraphs[0].text = "危险程度"
					cell.paragraphs[0].style = style 
					count_account_vul = 0 #统计已写入的漏洞数量数量
					vul_list = []
					df1 = pd.DataFrame(columns=['A', 'B', 'C', 'D']) 
					for file in glob.glob(os.path.join(file_name, 'Linux账号风险_风险视图列表*.csv')):
						df1 = CSVRead(file)
					if df1.empty: #如果没有账号风险,删除这张表
						logging.info("没有账号风险")
						no_data_tables.append(table[q])#将没有数据的表存入数组，后续一并删除
						no_data_tables.append(table[q+1])
						no_data_paragraphs.append("$账号风险检查情况")#对应的说明字段也存入数组，后续一并删除
						break
					#选择危急或高危的漏洞
					for index, row_vul_list in df1.iterrows():
						if row_vul_list["风险项名"] not in vul_list:
							if row_vul_list["危险程度"] == "危急":
								list1 =[row_vul_list["危险程度"], row_vul_list["风险项名"], str(df1.value_counts("风险项名")[row_vul_list["风险项名"]])+"台"]
								table[q] = table_add_a_line(table[q], list1)
								list2 = [row_vul_list["风险项名"],str(df1.value_counts("风险项名")[row_vul_list["风险项名"]])+"台",row_vul_list["风险描述"]+"\r【修复建议】"+row_vul_list["修复建议"]+"\r\r【修复影响】"+row_vul_list["修复影响"]]
								table[q+1] = table_add_a_line(table[q+1], list2)
								count_account_vul = count_account_vul+1
								vul_list.append(row_vul_list["风险项名"])
								if count_account_vul > 9:
									break
							if row_vul_list["危险程度"] == "高危":
								list1 =[row_vul_list["危险程度"], row_vul_list["风险项名"], str(df1.value_counts("风险项名")[row_vul_list["风险项名"]])+"台"]
								table[q] = table_add_a_line(table[q], list1)
								list2 = [row_vul_list["风险项名"],str(df1.value_counts("风险项名")[row_vul_list["风险项名"]])+"台",row_vul_list["风险描述"]+"\r【修复建议】"+row_vul_list["修复建议"]+"\r\r【修复影响】"+row_vul_list["修复影响"]]
								table[q+1] = table_add_a_line(table[q+1], list2)
								count_account_vul = count_account_vul+1
								vul_list.append(row_vul_list["风险项名"])
								if count_account_vul > 9:
									break
					table[q]._tbl.remove(table[q].rows[len(table[q].rows)-1]._tr)#删除最后一行（最后一行是空的）
					table[q+1]._tbl.remove(table[q+1].rows[len(table[q+1].rows)-1]._tr)#删除最后一行（最后一行是空的）
	
	#输出各漏洞的统计情况
	for file in glob.glob(os.path.join(file_name, 'Linux弱密码_全部风险项列表*.csv')):
		df1 = CSVRead(file)	
	for file in glob.glob(os.path.join(file_name, 'windows弱密码*.csv')):
		df2 = CSVRead(file)		
	replace_dict2 = {
	"$vul_count":str(count_vul),
	"$weak_passwd_count":str(len(df1)+len(df2)),
	"$app_vul_count":str(count_app_vul),
	"$os_vul_count":str(count_os_vul),
	"$account _vul_count":str(count_account_vul),
	"$MySQL弱密码问题列表":"",
	"$SSH弱密码问题列表":"",
	"$Redis弱密码问题列表":"",
	"$Tomcat弱密码问题列表":"",
	"$Rsync弱密码问题列表":"",
	"$Weblogic弱密码问题列表":"",
	"$Jenkins弱密码问题列表":"",
	"$OpenLDAP弱密码问题列表":"",
	"$InfluxDB弱密码问题列表":"",
	"$ProFTPD弱密码问题列表":"",
	"$SVN弱密码问题列表":"",
	"$PPTP弱密码问题列表":"",
	"$VNC弱密码问题列表":"",
	"$OpenVPN弱密码问题列表":"",
	"表$vsftpd弱密码问题列表":"",
	"$windows系统账户弱密码问题列表":"",
	"$应用风险检查情况":"",
	"$系统风险检查情况":"",
	"$账号风险检查情况":""	
	}
	
	replace_dict3 = {"nan":""}
		
	#清理无用数据
	for x in no_data_tables:
		# ~ print(table[x])
		# ~ print(x)
		x._element.getparent().remove(x._element)
	remove_paragraphs(doc, no_data_paragraphs)
	logging.info("")
	logging.info("已删除无数据模块"+str(no_data_paragraphs))

	remove_table_paragraphs(doc, ["【检测原理】"]) #删除会使得表格过于冗余的内容
	doc = check_and_change(doc, replace_dict2)
	doc = check_and_change2(doc, replace_dict3)
	logging.info("巡检结束")
	logging.info(replace_dict)
	logging.info("=====================================================================")
	return(doc)

	# ~ response = HttpResponse(f, content_type='application/x-zip')
	# ~ response['Content-Disposition'] = 'attachment; filename="%s"' % filename
	

if __name__ == '__main__':
	logging.basicConfig(level=logging.INFO,filename='log.log',format="%(asctime)s:%(levelname)s:%(message)s")
	logging.info(f'----------------------程序启动----------------------')
	run(host="0.0.0.0", port=8080, reloader=True)  # reloader设置为True可以在更新代码时自动重载

